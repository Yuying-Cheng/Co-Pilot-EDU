"""Format generated tasks as teacher-facing classroom task sheets."""
from __future__ import annotations

from typing import Any, Dict, List


def _join(values: List[str], default: str = "无") -> str:
    return "、".join([str(v) for v in values if v]) or default


def format_task_sheet(task_data: Dict[str, Any]) -> str:
    chapter_title = task_data.get("chapter_title") or task_data.get("chapter_id", "未命名章节")
    tasks = task_data.get("tasks", [])
    default_rounds = max([int(t.get("interaction_requirements", {}).get("min_rounds", 10) or 10) for t in tasks] or [10])
    lines: List[str] = []
    lines.append(f"{chapter_title} · 课堂探究任务")
    lines.append("")
    lines.append(
        f"任务说明：每个任务需与大模型交互{default_rounds}次以上，每次交互应围绕上一轮回答继续推进，"
        "交互过程中必须至少包含三种不同的交互方式（如：询问、表达见解、审辨、猜想、想象、创新、回答苏格拉底式提问等）。"
        "提交完整对话记录，并根据任务要求提交相应成果，附个人学习心得。"
    )
    lines.append("")
    for index, task in enumerate(tasks, start=1):
        req = task.get("interaction_requirements", {})
        out = task.get("output_requirements", {})
        lines.append(f"任务{_cn_num(index)}：{task.get('title', f'任务{index}')}")
        lines.append(task.get("description", "").strip())
        inquiry = task.get("inquiry_points") or task.get("exploration_points") or []
        if inquiry:
            lines.append("你需要探究的问题：")
            for item in inquiry:
                lines.append(f"- {item}")
        hints = task.get("hints") or task.get("teacher_hint") or ""
        if hints:
            if isinstance(hints, list):
                lines.append("提示：" + "；".join(str(item) for item in hints))
            else:
                lines.append(f"提示：{hints}")
        lines.append(f"交互方式要求：至少包含{_join(req.get('required_types', []))}等方式；不少于{req.get('min_rounds', default_rounds)}轮。")
        word_limit = out.get("word_limit", "500-800字")
        outputs = _join(out.get("required_outputs", []), "探究报告、交互记录、学习反思")
        lines.append(f"成果要求：提交{outputs}；建议字数{word_limit}。")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _cn_num(value: int) -> str:
    nums = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
    if value <= 10:
        return nums[value]
    return str(value)


def export_task_sheet_docx(task_data: Dict[str, Any], path: str) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        _export_minimal_docx(format_task_sheet(task_data), path)
        return

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    for block in format_task_sheet(task_data).split("\n"):
        if not block.strip():
            doc.add_paragraph("")
        elif block.endswith("课堂探究任务"):
            p = doc.add_paragraph()
            run = p.add_run(block)
            run.bold = True
            run.font.size = Pt(18)
        elif block.startswith("任务") and "：" in block:
            p = doc.add_paragraph()
            run = p.add_run(block)
            run.bold = True
            run.font.size = Pt(13)
        else:
            doc.add_paragraph(block)
    doc.save(path)


def _export_minimal_docx(text: str, path: str) -> None:
    import html
    import zipfile

    paragraphs = []
    for line in text.splitlines():
        escaped = html.escape(line)
        if not escaped:
            paragraphs.append("<w:p/>")
        else:
            paragraphs.append(f"<w:p><w:r><w:t>{escaped}</w:t></w:r></w:p>")
    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {''.join(paragraphs)}
    <w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>
  </w:body>
</w:document>"""
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as docx:
        docx.writestr("[Content_Types].xml", content_types)
        docx.writestr("_rels/.rels", rels)
        docx.writestr("word/document.xml", document_xml)
